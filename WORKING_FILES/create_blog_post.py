from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Set color to blue
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    # Set underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Welcome to Doctrine & Data', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Author
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
author_run = author.add_run('By Thomas Reichert')
author_run.italic = True

doc.add_paragraph()

# Intro
p = doc.add_paragraph()
p.add_run('Welcome to ').font.size = Pt(11)
run = p.add_run('Doctrine & Data')
run.italic = True
run.font.size = Pt(11)
p.add_run('—a blog about intellectual property law through an empirical lens.').font.size = Pt(11)

p = doc.add_paragraph()
p.add_run("I'm Thomas Reichert, an IP attorney and Assistant Professor of Law at Southern Illinois University Simmons School of Law. My work sits at the intersection of legal doctrine and empirical analysis, asking a simple question: ")
run = p.add_run("Does the law work the way we think it does?")
run.italic = True

# Section: What This Blog Is About
doc.add_heading('What This Blog Is About', level=1)

doc.add_paragraph("Legal doctrine is full of multifactor tests, balancing frameworks, and totality-of-the-circumstances analyses. Courts and practitioners treat these frameworks as if every factor matters, as if the careful weighing of numerous considerations is what produces just outcomes.")

doc.add_paragraph("But what if that's not actually how decisions get made?")

p = doc.add_paragraph()
p.add_run("My recent research on trademark confusion doctrine—analyzing thousands of TTAB decisions—found that the traditional thirteen-factor ")
run = p.add_run("DuPont")
run.italic = True
p.add_run(" test effectively collapses to just two factors. Mark similarity and goods relatedness predict outcomes with over 99% accuracy. The other eleven factors? Largely noise.")

doc.add_paragraph("This blog will explore questions like these across intellectual property law and beyond:")

# Bullet points
bullets = [
    ("Trademark doctrine", "What do the data tell us about how confusion is actually decided?"),
    ("Patent practice", "Insights from prosecution, examination, and litigation"),
    ("Empirical legal studies", "Methodology, findings, and implications for practice"),
    ("Law and technology", "Including the growing role of AI in legal analysis and practice"),
]

for bold_text, rest in bullets:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold_text + ": ")
    run.bold = True
    p.add_run(rest)

# Section: Who This Is For
doc.add_heading('Who This Is For', level=1)

doc.add_paragraph("I write for practitioners who want actionable insights, academics interested in empirical approaches to doctrine, and anyone curious about what the data reveal when we look closely at how law actually operates.")

doc.add_paragraph("The tone here will be rigorous but (I hope) readable. Complex ideas deserve clear explanations, and empirical findings should translate into practical understanding.")

# Section: What's Next
doc.add_heading("What's Next", level=1)

p = doc.add_paragraph()
p.add_run("In upcoming posts, I'll dig deeper into the ")
run = p.add_run("DuPont")
run.italic = True
p.add_run(" findings and their implications for trademark practice, explore what similar methods might reveal in other areas of IP, and comment on developments in the field as they arise.")

# Links paragraph
p = doc.add_paragraph("If you'd like to follow along, you can connect with me on ")
add_hyperlink(p, "LinkedIn", "https://www.linkedin.com/in/thomas-reichert-071376173/")
p.add_run(".")

doc.add_paragraph("Thanks for reading.")

# Signature
p = doc.add_paragraph()
p.add_run("—TR").font.size = Pt(11)

doc.add_paragraph()

# Divider
doc.add_paragraph("—" * 30)

doc.add_paragraph()

# Learn More section
learn_more = doc.add_heading('Learn More', level=2)

p = doc.add_paragraph("Visit ")
add_hyperlink(p, "reichert.law", "https://reichert.law")
p.add_run(" to learn more about my academic research, consulting services, and publications.")

p = doc.add_paragraph("Read the full paper: ")
add_hyperlink(p, "Doctrine, Data, and the Death of DuPont", "https://papers.ssrn.com/abstract=5843642")
p.add_run(" on SSRN.")

p = doc.add_paragraph("Browse the ")
add_hyperlink(p, "blog archive", "https://doctrineanddata.blogspot.com/")
p.add_run(" for previous posts.")

# Save
doc.save(r'E:\GitHub\TTAB-Prediction-Tool-Web\Reichert.law\WORKING_FILES\Blog_Post_Welcome.docx')
print("Document created successfully!")
